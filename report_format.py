from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, Inches, Cm


def set_cell_border(cell, **kwargs):
    """
    изменяет параметры таблицы,
    использование:
    
    set_cell_border(
        cell,
        top={"sz": 12, "val": "single", "color": "#FF0000", "space": "0"},
        bottom={"sz": 12, "color": "#00FF00", "val": "single"},
        start={"sz": 24, "val": "dashed", "shadow": "true"},
        end={"sz": 12, "val": "dashed"},
    )
    """
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()

    # check for tag existnace, if none found, then create one
    tcBorders = tcPr.first_child_found_in("w:tcBorders")
    if tcBorders is None:
        tcBorders = OxmlElement('w:tcBorders')
        tcPr.append(tcBorders)

    # list over all available tags
    for edge in ('start', 'top', 'end', 'bottom', 'insideH', 'insideV'):
        edge_data = kwargs.get(edge)
        if edge_data:
            tag = 'w:{}'.format(edge)

            # check for tag existnace, if none found, then create one
            element = tcBorders.find(qn(tag))
            if element is None:
                element = OxmlElement(tag)
                tcBorders.append(element)

            # looks like order of attributes is important
            for key in ["sz", "val", "color", "space", "shadow"]:
                if key in edge_data:
                    element.set(qn('w:{}'.format(key)), str(edge_data[key]))

# создал docx
document = Document()

# отредактировал отступы от краев листа
sections = document.sections
for section in sections:
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(1.0)

# создал таблицу 3х2
sdl_table = document.add_table(rows=3, cols=2)

# отредактировал ширину колонок
column0 = sdl_table.columns[0]
column0.width = Inches(2.1)
# мб будет проще счетчика, если получится через for перебрать все раны параграфов
# в первом столбе
column1 = sdl_table.columns[1]
column1.width = Inches(5)

# заполнение ячеек
cell0 = sdl_table.cell(row_idx=0, col_idx=0).text = "Источник:"
cell1 = sdl_table.cell(row_idx=0, col_idx=1).text = "«nrc.gov»"
cell2 = sdl_table.cell(row_idx=1, col_idx=0).text = "Опубликовано:"
cell3 = sdl_table.cell(row_idx=1, col_idx=1).text = "03.07.2024"
cell4 = sdl_table.cell(row_idx=2, col_idx=0).text = "Ссылка на источник:"
cell5 = sdl_table.cell(row_idx=2, col_idx=1).text = "https://www.nrc.gov/reading-rm/doc-collections/event-status/event/2024/20240711en.html#en57208"

# обнуление обводки, смена шрифта, размера шрифта + bold
# тут счетчик, самая простая реализация выделения первого столбца
counter = 0
for row in sdl_table.rows:
    for cell in row.cells:
        cell.paragraphs[0].add_run('')
        rc = cell.paragraphs[0].runs[0]
        rc.font.size = Pt(14)
        rc.font.name = 'TimesNewRoman'
        if counter == 0:
            rc.font.bold = True
            rc.font.underline = True
            counter += 1
        else:
            counter = 0
        set_cell_border(
            cell,
            top={"sz": 0},
            bottom={"sz": 0},
            start={"sz": 0},
            end={"sz": 0},
        )


document.add_paragraph()
art_header = document.add_paragraph('')
art_header.paragraph_format.first_line_indent = Inches(0.5)
art_header.add_run('В США, штат Миссури, город Осейдж-Бич был поврежден плотномер, содержащий цезий-137 и амерций-241')
art_header.runs[0].font.name = 'TimesNewRoman'
art_header.runs[0].font.size = Pt(14)
art_header.runs[0].font.bold = True

article = document.add_paragraph('')
article.paragraph_format.first_line_indent = Inches(0.5)
article.add_run('В 11:46 CDT 7/3/2024 сотрудник по радиационной безопасности в Midwest Subsurface Testing сообщил, что на строительной площадке был поврежден датчик. Датчик плотности влаги InstroTek MC1 Elite, содержащий 10 милликюри цезия-137 и 50 милликюри америция-241/бериллия, был опрокинут погрузчиком. Источник застрял в экранированном положении. Было проведено радиологическое обследование, которое подтвердило отсутствие загрязнения. Поврежденный датчик был извлечен и доставлен на предприятие поставщика для проведения испытания на герметичность.\nЭто событие было сообщено в соответствии с 10 CFR 30.50 (b)(2).')
article.runs[0].font.name = 'TimesNewRoman'
article.runs[0].font.size = Pt(14)

document.save('demo.docx')
