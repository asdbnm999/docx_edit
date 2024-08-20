import os
import shutil
from docx import Document
from docx.shared import Pt, Cm
from SINGLE_get_info_from_txt_refractored import get_docx_info
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT


def detect_pdfiles():           # :DDD
    current_dir = os.getcwd()
    tables_dir = os.path.join(current_dir, 'tables')
    batches_dir = os.path.join(current_dir, 'batches')

    os.makedirs(batches_dir, exist_ok=True)

    pdf_files = [f for f in os.listdir(tables_dir) if f.endswith('.pdf')]

    for file in pdf_files:
        if 'Batch' in file:
            print(f"Ошибка: файл '{file}' содержит множественный отчет. \nПеремещение в папку 'batches'.")
            shutil.move(os.path.join(tables_dir, file), os.path.join(batches_dir, file))

    remaining_files = [f for f in os.listdir(tables_dir) if f.endswith('.pdf')]
    return remaining_files


def create_report_table():
    pdfiles_list = detect_pdfiles()
    pdfiles_list.sort()
    print(pdfiles_list)

    document = Document()

    section = document.sections[0]
    section.orientation = 1  # 1 - горизонтальная, 0 - вертикальная
    section.page_width = section.page_height
    section.page_height = section.page_width

    # отредактировал отступы от краев листа
    sections = document.sections
    for section in sections:
        section.orientation = 1
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(1)
        section.right_margin = Cm(1.0)

    report_table = document.add_table(rows=(len(pdfiles_list) + 2), cols=8)

    report_table.columns[0].width = Cm(1.5)
    report_table.columns[1].width = Cm(3.0)
    report_table.columns[2].width = Cm(2.3)
    report_table.columns[3].width = Cm(3.0)
    report_table.columns[4].width = Cm(3.5)
    report_table.columns[5].width = Cm(4.5)
    report_table.columns[6].width = Cm(4.5)
    report_table.columns[7].width = Cm(6.0)

    cell0 = report_table.cell(row_idx=0, col_idx=0).text = "№\nп/п"
    cell1 = report_table.cell(row_idx=0, col_idx=1).text = "№ из базы\nданных\nМАГАТЭ"
    cell2 = report_table.cell(row_idx=0, col_idx=2).text = "Дата обнаружения"
    cell0 = report_table.cell(row_idx=0, col_idx=3).text = "Страна"
    cell1 = report_table.cell(row_idx=0, col_idx=4).text = "Место/Регион"
    cell2 = report_table.cell(row_idx=0, col_idx=5).text = "Что обнаружено"
    cell0 = report_table.cell(row_idx=0, col_idx=6).text = "Инцидент с ИИИ"
    cell1 = report_table.cell(row_idx=0, col_idx=7).text = "Описание факта обнаружения ИИИ"

    for cell in report_table.rows[0].cells:

        # Выравнивание текста по центру (горизонтально)
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        # Выравнивание текста по центру (вертикально)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

    for j in range(len(pdfiles_list)):
        table_info = get_docx_info(pdfiles_list[j])
        for i, cell in enumerate(report_table.rows[j+1].cells):
            if i != 0 and i != 1 and i != 3 and i != 4 and i != 6:
                cell.text = table_info[i - 1]
                cell_paragraph = cell.paragraphs[0]
                if i == 7:
                    cell_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
                else:
                    cell_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                run = cell_paragraph.runs[0]
                run.font.size = Pt(12)

            elif i == 0:
                cell.text = str(j+1)
                cell_paragraph = cell.paragraphs[0]
                cell_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                run = cell_paragraph.runs[0]
                run.font.size = Pt(12)

            elif i == 1:
                cell.text = table_info[i - 1]
                cell_paragraph = cell.paragraphs[0]
                cell_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                run1 = cell_paragraph.runs[0]
                run1.font.size = Pt(12)
                run2 = cell_paragraph.add_run("\n" + table_info[-4])
                run2.font.size = Pt(10)

            elif i == 3:  # Для четвертого столбца
                cell.text = table_info[i - 1]
                cell_paragraph = cell.paragraphs[0]
                cell_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                run1 = cell_paragraph.runs[0]
                run1.font.size = Pt(12)
                run2 = cell_paragraph.add_run("\n" + table_info[-3])
                run2.font.size = Pt(10)

            elif i == 4:  # Для четвертого столбца
                cell.text = table_info[i - 1]
                cell_paragraph = cell.paragraphs[0]
                cell_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                run1 = cell_paragraph.runs[0]
                run1.font.size = Pt(12)
                run2 = cell_paragraph.add_run("\n" + table_info[-2])
                run2.font.size = Pt(10)

            elif i == 6:  # Для четвертого столбца
                cell.text = table_info[i - 1]
                cell_paragraph = cell.paragraphs[0]
                cell_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                run1 = cell_paragraph.runs[0]
                run1.font.size = Pt(12)
                run2 = cell_paragraph.add_run("\n" + table_info[-1])
                run2.font.size = Pt(10)

        for row in report_table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    # Устанавливаем шрифт и размер
                    run = paragraph.runs
                    if run:
                        for r in run:
                            r.font.name = 'Times New Roman'

        document.save('demo.docx')


create_report_table()
