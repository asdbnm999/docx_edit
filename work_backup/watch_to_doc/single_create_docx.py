from docx import Document
from docx.shared import Pt, Cm
from SINGLE_get_info_from_txt_refractored import get_docx_info
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

table_info = get_docx_info('2024-05-002.pdf')

document = Document()

section = document.sections[0]
section.orientation = 1  # 1 - горизонтальная, 0 - вертикальная
section.page_width = section.page_height
section.page_height = section.page_width

# отредактировал отступы от краев листа
sections = document.sections
for section in sections:
    section.orientation = 1
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(1)
    section.right_margin = Cm(1.0)

# создал таблицу 3х2
report_table = document.add_table(rows=2, cols=8)

# отредактировал ширину колонок
report_table.columns[0].width = Cm(1.5)
report_table.columns[1].width = Cm(3.0)
report_table.columns[2].width = Cm(2.3)
report_table.columns[3].width = Cm(3.0)
report_table.columns[4].width = Cm(3.5)
report_table.columns[5].width = Cm(4.5)
report_table.columns[6].width = Cm(4.5)
report_table.columns[7].width = Cm(8.0)



for i, cell in enumerate(report_table.rows[0].cells):
    if i != 0 and i != 1 and i != 3 and i != 4 and i != 6:
        cell.text = table_info[i-1]
        cell_paragraph = cell.paragraphs[0]
        if i == 7:
            cell_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
        else:
            cell_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        run = cell_paragraph.runs[0]
        run.font.size = Pt(12)
        run.font.name = 'TimesNewRoman'

    elif i == 0:
        cell.text = '1'
        cell_paragraph = cell.paragraphs[0]
        cell_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        run = cell_paragraph.runs[0]
        run.font.size = Pt(12)
        run.font.name = 'TimesNewRoman'

    elif i == 1:
        cell.text = table_info[i-1]
        cell_paragraph = cell.paragraphs[0]
        cell_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        run1 = cell_paragraph.runs[0]
        run1.font.size = Pt(12)
        run1.font.name = 'TimesNewRoman'
        run2 = cell_paragraph.add_run("\n" + table_info[-4])
        run2.font.size = Pt(10)
        run2.font.name = 'TimesNewRoman'

    elif i == 3:  # Для четвертого столбца
        cell.text = table_info[i-1]
        cell_paragraph = cell.paragraphs[0]
        cell_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        run1 = cell_paragraph.runs[0]
        run1.font.size = Pt(12)
        run1.font.name = 'TimesNewRoman'
        run2 = cell_paragraph.add_run("\n" + table_info[-3])
        run2.font.size = Pt(10)
        run2.font.name = 'TimesNewRoman'

    elif i == 4:  # Для четвертого столбца
        cell.text = table_info[i-1]
        cell_paragraph = cell.paragraphs[0]
        cell_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        run1 = cell_paragraph.runs[0]
        run1.font.size = Pt(12)
        run1.font.name = 'TimesNewRoman'
        run2 = cell_paragraph.add_run("\n" + table_info[-2])
        run2.font.size = Pt(10)
        run2.font.name = 'TimesNewRoman'

    elif i == 6:  # Для четвертого столбца
        cell.text = table_info[i-1]
        cell_paragraph = cell.paragraphs[0]
        cell_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        run1 = cell_paragraph.runs[0]
        run1.font.size = Pt(12)
        run1.font.name = 'TimesNewRoman'
        run2 = cell_paragraph.add_run("\n" + table_info[-1])
        run2.font.size = Pt(10)
        run2.font.name = 'TimesNewRoman'


document.save('demo.docx')
