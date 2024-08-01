import tkinter as tk
from tkinter import messagebox as mb
from tkinter import ttk
from tkinter import filedialog
import webbrowser as wb
from deep_translator import GoogleTranslator
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, Inches, Cm


class MenuFrame(tk.Frame):
    def __init__(self, parent):
        super().__init__(parent)
        #############################
        # поменять названия на теги #
        #############################
        menu_canvas = tk.Canvas(self, width=330, height=185, bg='#4B0082')
        monitoring_but = menu_canvas.create_text((10, 22), anchor='nw',
                                                 text=f'Мониторинг',
                                                 font='Consolas 14',
                                                 fill='#E0FFFF',
                                                 activefill='#00BFFF')
        menu_canvas.tag_bind(monitoring_but, "<Button-1>", lambda win='ws': goto(event=None, win='ws'))

        panel = menu_canvas.create_image((120, 22), anchor='nw', image=panel_img)
        menu_canvas.tag_bind(panel, "<Button-1>", lambda win='ws': goto(event=None, win='ws'))

        docx_reports_compile_but = menu_canvas.create_text((10, 80), anchor='nw',
                                                           text=f'Составление отчетов .docx',
                                                           font='Consolas 14',
                                                           fill='#E0FFFF',
                                                           activefill='#00BFFF')
        menu_canvas.tag_bind(docx_reports_compile_but, "<Button-1>", lambda win='cr': goto(event=None, win='cr'))

        report = menu_canvas.create_image((267, 79), anchor='nw', image=report_img)
        menu_canvas.tag_bind(report, "<Button-1>", lambda win='ws': goto(event=None, win='cr'))

        translate_tables_but = menu_canvas.create_text((10, 145), anchor='nw',
                                                       text=f'Перевод таблиц БД МАГАТЭ',
                                                       font='Consolas 14',
                                                       fill='#E0FFFF',
                                                       activefill='#00BFFF')
        menu_canvas.tag_bind(translate_tables_but, "<Button-1>", lambda win='tt': goto(event=None, win='tt'))

        table = menu_canvas.create_image((257, 144), anchor='nw', image=table_img)
        menu_canvas.tag_bind(table, "<Button-1>", lambda win='ws': goto(event=None, win='tt'))

        settings_but = menu_canvas.create_image((306, 5), anchor='nw', image=param_img)
        menu_canvas.tag_bind(settings_but, "<Button-1>", self.open_settings_window)

        info_but = menu_canvas.create_image((280, 5), anchor='nw', image=info_img)
        menu_canvas.tag_bind(info_but, '<Button-1>', self.show_info)

        restore_settings_but = menu_canvas.create_image((306, 30), anchor='nw', image=restore_img)
        menu_canvas.tag_bind(restore_settings_but, "<Button-1>", self.restore_settings)

        menu_canvas.pack()
        self.pack()

    def show_info(self, event):
        print('asdasdasdasdasdasdasdasd')

    def restore_settings(self, event):
        restore_settings_flag = mb.askyesno(title='Сброс настроек.', message='Подтвердите действие.')
        if restore_settings_flag:
            with open('config.txt', 'w') as file:
                file.write('browser = default' + '\nbrowser_path = default' + '\ndownload_path = default')
            set_config()

            mb.showinfo(title='Отчет.',
                        message=f'Настройки сброшены.')
        else:
            pass

    def open_settings_window(self, event):
        self.sw = tk.Toplevel(self, width=300, height=150)

        self.sw.resizable(False, False)
        self.sw_canvas = tk.Canvas(self.sw, width=300, height=150, bg='#4B0082')

        self.select_setting_lab = self.sw_canvas.create_text((10, 10),
                                                             anchor='nw',
                                                             text='Выберите пункт для настройки',
                                                             fill='#E0FFFF',
                                                             font='Consolas 10',
                                                             )
        settings = ['Браузер', 'Путь сохранения файлов']
        self.select_setting = ttk.Combobox(self.sw,
                                           values=settings,
                                           state='readonly',
                                           width=25,
                                           justify=tk.CENTER,
                                           )
        self.select_setting.place(x=10, y=30)
        self.select_setting.bind('<<ComboboxSelected>>', self.selected1)

        self.sw_canvas.pack()

    def selected1(self, event):
        self.set_choice = self.select_setting.get()
        self.sw_canvas.create_text((10, 53),
                                   anchor='nw',
                                   text='',
                                   fill='#E0FFFF',
                                   tags='second_annotation',
                                   font='Consolas 10')
        if self.set_choice == 'Браузер':
            try:
                self.sw_canvas.coords('get_path_but', (-100, -100))
                self.sw_canvas.coords('confirm_but', (-100, -100))
                self.sw_canvas.coords('dismiss_but', (-100, -100))
            except Exception as e:
                pass

            self.sw_canvas.itemconfig('second_annotation', text='Выберите браузер')

            browsers = ['Opera', 'Chrome', 'Chromium', 'FireFox']
            self.select_browser = ttk.Combobox(self.sw,
                                               values=browsers,
                                               state='readonly',
                                               width=25,
                                               justify=tk.CENTER
                                               )
            self.select_browser.place(x=10, y=70)
            self.select_browser.bind('<<ComboboxSelected>>', self.selected_browser)

        elif self.set_choice == 'Путь сохранения файлов':
            try:
                self.select_browser.place(x=-100, y=-100)
                self.sw_canvas.coords('get_BPath_but', (-100, -100))
                self.sw_canvas.coords('confirm_but', (-100, -100))
                self.sw_canvas.coords('dismiss_but', (-100, -100))
            except Exception as e:
                pass

            self.sw_canvas.itemconfig('second_annotation', text='Выберите директорию')

            try:
                self.hand_ent_path.place(x=-100, y=-100)
            except:
                pass

            self.hand_ent_path = tk.Entry(self.sw, width=27, bg='#6A5ACD', fg='#E0FFFF')
            self.hand_ent_path.place(x=10, y=70)

            self.sw_canvas.create_text((180, 72),
                                       anchor='nw',
                                       text='Обзор...',
                                       fill='lightblue',
                                       activefill='white',
                                       font='Consolas 10',
                                       tags='get_path_but')
            self.sw_canvas.tag_bind('get_path_but', '<Button-1>', self.select_directory)

            self.sw_canvas.create_image((245, 125), anchor='nw', image=confirm_img, tags='confirm_but')
            self.sw_canvas.tag_bind('confirm_but', '<Button-1>', self.confirm_changes)

            self.sw_canvas.create_image((275, 125), anchor='nw', image=dismiss_img, tags='dismiss_but')
            self.sw_canvas.tag_bind('dismiss_but', '<Button-1>', self.dismiss_changes)

    def selected_browser(self, event):
        self.browser_name = self.select_browser.get().lower()
        print(self.browser_name)

        try:
            self.hand_ent_path.place(x=-100, y=-100)
        except:
            pass

        self.hand_ent_path = tk.Entry(self.sw, width=27, bg='#6A5ACD', fg='#E0FFFF')
        self.hand_ent_path.place(x=10, y=110)

        self.sw_canvas.create_text((180, 112),
                                   anchor='nw',
                                   text='Обзор...',
                                   fill='lightblue',
                                   activefill='white',
                                   font='Consolas 10',
                                   tags='get_BPath_but')
        self.sw_canvas.tag_bind('get_BPath_but', '<Button-1>', self.select_launcher)

        self.sw_canvas.create_image((245, 125), anchor='nw', image=confirm_img, tags='confirm_but')
        self.sw_canvas.tag_bind('confirm_but', '<Button-1>', self.confirm_changes)

        self.sw_canvas.create_image((275, 125), anchor='nw', image=dismiss_img, tags='dismiss_but')
        self.sw_canvas.tag_bind('dismiss_but', '<Button-1>', self.dismiss_changes)

    def select_directory(self, event):
        self.directory = filedialog.askdirectory()
        self.hand_ent_path.insert(0, self.directory)
        self.sw.lift()

    def select_launcher(self, event):
        self.file_path = filedialog.askopenfile().name
        self.hand_ent_path.insert(0, self.file_path)
        self.sw.lift()

    def confirm_changes(self, event):
        """
        browser = default
        browser_path = default
        download_path = default
        C:/Users/vboxuser/AppData/Local/Programs/Opera/launcher.exe
        """
        global browser, browser_path, download_path
        print(browser, browser_path, download_path)
        if self.set_choice == 'Браузер':
            try:
                browser = self.browser_name
                browser_path = self.hand_ent_path.get()
                if '\\' not in browser_path and ':/' in browser_path:
                    wb.register(self.browser_name,
                                None,
                                wb.BackgroundBrowser(browser_path))

                    with open('config.txt', 'w') as file:
                        file.write('browser = ' + browser + '\nbrowser_path = ' +
                                   browser_path + '\ndownload_path = ' + download_path)

                    mb.showinfo(title='Отчет.',
                                message=f'Параметры применены успешно!\nБраузер успешно сменен на {browser}')
                else:
                    mb.showwarning(title='Ошибка настройки.', message='Новые параметры не применены!\n'
                                                                      'Проверьте путь скачивания, в нем не должно быть'
                                                                      ' символов "\\"')
            except Exception as e:
                print(e)
                mb.showwarning(title='Ошибка настройки.', message='Новые параметры не применены!')

        elif self.set_choice == 'Путь сохранения файлов':
            try:
                download_path = self.hand_ent_path.get()
                if '\\' not in browser_path and ':/' in browser_path:

                    with open('config.txt', 'w') as file:
                        file.write('browser = ' + browser + '\nbrowser_path = ' +
                                   browser_path + '\ndownload_path = ' + download_path)

                    mb.showinfo(title='Отчет.', message='Параметры применены успешно!\nПуть для скачивания файлов:\n'
                    f'{download_path}')
                else:
                    mb.showwarning(title='Ошибка настройки.', message='Новые параметры не применены!\n'
                                                                      'Проверьте путь скачивания, в нем не должно быть'
                                                                      ' символов "\\"')
            except Exception as e:
                print(e)
                mb.showwarning(title='Ошибка настройки.', message='Новые параметры не применены!')

        else:
            mb.showwarning(title='Ошибка настройки.', message='Новые параметры не применены!')

    def dismiss_changes(self, event):
        self.sw.destroy()


class WebSourcesFrame(tk.Frame):
    def __init__(self, parent):
        super().__init__(parent)
        self.ws_canvas = tk.Canvas(self, width=800, height=600, bg='#6A5ACD')
        self.ws_canvas.create_image((10, 10), anchor='nw', image=back_img, tags='back_but')
        self.ws_canvas.tag_bind('back_but', '<Button-1>', lambda win='menu': goto(event=None, win='menu'))

        # размещение ссылок на canvas и добавление их в список
        link_items = []
        for i, link in enumerate(smi_list):
            if i <= 19:
                coords = (100, (i * 30) + 20)  # x и y координаты для каждой ссылки
            elif i > 19:
                i -= 20
                coords = (360, (i * 30) + 20)

            short_link = link[8:-1]
            if short_link == 'www.nrc.gov/reading-rm/doc-collections/event-status/event/2024/index.htm': short_link = 'nrc.gov'

            item = self.ws_canvas.create_text(*coords, text=short_link,
                                              fill="#E0FFFF",
                                              activefill='#00BFFF',
                                              anchor="w",
                                              font='Consolas 12',
                                              tags="link")
            link_items.append((item, link))

        # Обработчик клика
        def on_click(event):
            global browser
            clicked_item = self.ws_canvas.find_withtag("current")
            if clicked_item:
                for item, link in link_items:
                    self.search_kw = self.keyword_ent.get().strip()
                    if self.search_kw != '':
                        if link == 'https://russian.rt.com/' or link == 'https://regnum.ru/':
                            link += f'search?q={self.search_kw}'
                        elif link == 'https://rbc.ru/':
                            link += f'search/?query={self.search_kw}'
                        elif link == 'https://lenta.ru/' or link == 'https://vedomosti.ru/':
                            link += f'search?query={self.search_kw}'
                        elif link == 'https://rg.ru/' or link == 'https://iz.ru/' or link == 'https://aif.ru/' or link == 'https://tass.ru/' or link == 'https://moslenta.ru/':
                            link += f'search?text={self.search_kw}'
                        elif link == 'https://kommersant.ru/':
                            link += f'search/results?search_query={self.search_kw}'
                        elif link == 'https://gazeta.ru/':
                            link += f'search.shtml?text={self.search_kw}&p=main&input=utf8'
                        elif link == 'https://mk.ru/':
                            link += f'search/?q={self.search_kw}'
                        elif link == 'https://life.ru/':
                            link += f'search?search={self.search_kw}'
                        elif link == 'https://info.sibnet.ru/':
                            link += f'articles/?text={self.search_kw}'
                        elif link == 'https://news.yahoo.com/':
                            self.translate_search_kw('en')
                            link = f'https://search.yahoo.com/search?p={self.search_kw}&fr=yfp-t&fr2=p%3Afp%2Cm%3Asb&ei=UTF-8&fp=1'
                        elif link == 'https://bloomberg.com/' or link == 'https://cnbc.com/' or link == 'https://washingtonpost.com/':
                            self.translate_search_kw('en')
                            link += f'search/?query={self.search_kw}'
                        elif link == 'https://emm.newsbrief.eu/':
                            link += f'NewsBrief/dynamic?language=ru&page=1&edition=searchresults&option=&atLeast={self.search_kw}'
                        elif link == 'https://www.wsj.com' or link == 'https://cnn.com/':
                            self.translate_search_kw('en')
                            link += f'search?query={self.search_kw}'
                        elif link == 'https://nytimes.com/':
                            self.translate_search_kw('en')
                            link += f'search?dropmab=false&query={self.search_kw}&sort=newest'
                        elif link == 'https://foxnews.com/':
                            self.translate_search_kw('en')
                            link += f'search-results/search?q={self.search_kw}'
                        elif link == 'https://www.npr.org/':
                            self.translate_search_kw('en')
                            link += f'search/?query={self.search_kw}&page=1&sortType=byDate'
                        elif link == 'https://reuters.com/':
                            self.translate_search_kw('en')
                            link += f'site-search/?query={self.search_kw}'
                        elif link == 'https://usatoday.com/' or link == 'https://reddit.com/' or link == 'https://mirror.co.uk/':
                            self.translate_search_kw('en')
                            link += f'search/?q={self.search_kw}'
                        elif link == 'https://latimes.com/':
                            self.translate_search_kw('en')
                            link += f'search?q={self.search_kw}'
                        elif link == 'https://www.nrc.gov/reading-rm/doc-collections/event-status/event/2024/index.html':
                            self.translate_search_kw('en')
                            link = 'https://www.nrc.gov/site-help/search.html?site=AllSites&searchtext={self.search_kw}'
                        elif link == 'https://spiegel.de/':
                            self.translate_search_kw('de')
                            link += f'suche/?suchbegriff={self.search_kw}'
                        elif link == 'https://tagesschau.de/':
                            self.translate_search_kw('de')
                            link += f'suche#/all/1/?searchText={self.search_kw}'
                        elif link == 'https://zeit.de/':
                            self.translate_search_kw('de')
                            link += f'suche/index?q={self.search_kw}'
                        elif link == 'https://lemonde.fr/':
                            self.translate_search_kw('fr')
                            link += f'recherche/?search_keywords={self.search_kw}&search_sort=dateCreated_desc&page=1'
                        elif link == 'https://liberation.fr/' or link == 'https://leparisien.fr/':
                            self.translate_search_kw('fr')
                            link += f'recherche/?query={self.search_kw}'
                        elif link == 'https://theguardian.com/':
                            self.translate_search_kw('en')
                            link = f'https://www.google.co.uk/search?as_q={self.search_kw}&as_epq=&as_oq=&as_eq=&as_nlo=&as_nhi=&lr=&cr=&as_qdr=all&as_sitesearch=www.theguardian.com&as_occt=any&as_filetype=&tbs='
                        elif link == 'https://economist.com/':
                            self.translate_search_kw('en')
                            link += f'search?q={self.search_kw}&sort=date'
                        if browser == 'default':
                            wb.open_new_tab(link)
                        else:
                            wb.get(browser).open_new_tab(link)
                        break

        # Связывание события клика с обработчиком
        self.ws_canvas.tag_bind("link", "<Button-1>", on_click)

        # добавление русских флагов первой колонки
        for i in range(0, 20):
            flag_x = 70
            flag_y = (i * 30) + 10
            if (i >= 17) and (i <= 19):
                self.insert_country_flag(coords=(flag_x, flag_y), country='eu')
            else:
                self.insert_country_flag(coords=(flag_x, flag_y), country='ru')

        # добавление флагов второй колонки 330 10
        for i in range(0, 20):
            flag_x = 330
            flag_y = (i * 30) + 10
            if i >= 17:
                self.insert_country_flag(coords=(flag_x, flag_y), country='uk')
            elif (i >= 14) and (i <= 16):
                self.insert_country_flag(coords=(flag_x, flag_y), country='fr')
            elif (i >= 11) and (i <= 13):
                self.insert_country_flag(coords=(flag_x, flag_y), country='ge')
            else:
                self.insert_country_flag(coords=(flag_x, flag_y), country='us')

        self.keyword_ent = tk.Entry(self, width=30, bg='#6A5ACD', fg='#E0FFFF', font='Consolas 10')
        self.keyword_ent.place(x=570, y=19)

        self.search_kw = ''
        self.ws_canvas.create_image((540, 17), anchor='nw', image=restore_kw_img, tags='restore_kw')
        self.ws_canvas.tag_bind('restore_kw', '<Button-1>', self.restore_search_kw)

        self.ws_canvas.pack()
        self.pack()

    def translate_search_kw(self, target_lang: str):
        self.search_kw = GoogleTranslator(source='auto', target=target_lang).translate(self.search_kw)

    def restore_search_kw(self, event):
        self.keyword_ent.delete(0, tk.END)

    def insert_country_flag(self, coords: tuple, country: str):
        """
        добавление флагов стран для ссылок на СМИ

        :param coords: (x, y)
        :param country: 'ru', 'uk', 'us', 'eu', 'fr', 'ge'
        :return:
        """
        if country == 'ru':
            self.ws_canvas.create_image(coords, anchor='nw', image=ru_flag)
        elif country == 'uk':
            self.ws_canvas.create_image(coords, anchor='nw', image=uk_flag)
        elif country == 'us':
            self.ws_canvas.create_image(coords, anchor='nw', image=us_flag)
        elif country == 'fr':
            self.ws_canvas.create_image(coords, anchor='nw', image=fr_flag)
        elif country == 'ge':
            self.ws_canvas.create_image(coords, anchor='nw', image=ge_flag)
        elif country == 'eu':
            self.ws_canvas.create_image(coords, anchor='nw', image=eu_flag)


class CompileReportsFrame(tk.Frame):
    def __init__(self, parent):
        super().__init__(parent)
        self.cr_canvas = tk.Canvas(self, width=800, height=600, bg='#6A5ACD')
        self.cr_canvas.create_image((10, 10), anchor='nw', image=back_img, tags='back_but')
        self.cr_canvas.tag_bind('back_but', '<Button-1>', lambda win='menu': goto(event=None, win='menu'))
        self.cr_canvas.pack()
        self.pack()

    def compile_report(self):
        def set_cell_border(cell, **kwargs):
            """
            изменяет параметры таблицы,
            использование:

            set_cell_border(
                cell,
                top={"sz": 12, "val": "single", "color": "#FF0000", "space": "0"},
                bottom={"sz": 12, "color": "#00FF00", "val": "single"},
                start={"sz": 24, "val": "dashed", "shadow": "true"},
                end={"sz": 12, "val": "dashed"},
            )
            """
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()

            # check for tag existnace, if none found, then create one
            tcBorders = tcPr.first_child_found_in("w:tcBorders")
            if tcBorders is None:
                tcBorders = OxmlElement('w:tcBorders')
                tcPr.append(tcBorders)

            # list over all available tags
            for edge in ('start', 'top', 'end', 'bottom', 'insideH', 'insideV'):
                edge_data = kwargs.get(edge)
                if edge_data:
                    tag = 'w:{}'.format(edge)

                    # check for tag existnace, if none found, then create one
                    element = tcBorders.find(qn(tag))
                    if element is None:
                        element = OxmlElement(tag)
                        tcBorders.append(element)

                    # looks like order of attributes is important
                    for key in ["sz", "val", "color", "space", "shadow"]:
                        if key in edge_data:
                            element.set(qn('w:{}'.format(key)), str(edge_data[key]))

        document = Document()

        # отредактировал отступы от краев листа
        sections = document.sections
        for section in sections:
            section.top_margin = Cm(2.0)
            section.bottom_margin = Cm(2.0)
            section.left_margin = Cm(2.5)
            section.right_margin = Cm(1.0)

        # создал таблицу 3х2
        sdl_table = document.add_table(rows=3, cols=2)

        # отредактировал ширину колонок
        column0 = sdl_table.columns[0]
        column0.width = Inches(2.1)
        # мб будет проще счетчика, если получится через for перебрать все раны параграфов
        # в первом столбе
        column1 = sdl_table.columns[1]
        column1.width = Inches(5)

        # заполнение ячеек
        cell0 = sdl_table.cell(row_idx=0, col_idx=0).text = "Источник:"
        cell1 = sdl_table.cell(row_idx=0, col_idx=1).text = "«nrc.gov»"
        cell2 = sdl_table.cell(row_idx=1, col_idx=0).text = "Опубликовано:"
        cell3 = sdl_table.cell(row_idx=1, col_idx=1).text = "03.07.2024"
        cell4 = sdl_table.cell(row_idx=2, col_idx=0).text = "Ссылка на источник:"
        cell5 = sdl_table.cell(row_idx=2,
                               col_idx=1).text = "https://www.nrc.gov/reading-rm/doc-collections/event-status/event/2024/20240711en.html#en57208"

        # обнуление обводки, смена шрифта, размера шрифта + bold
        # тут счетчик, самая простая реализация выделения первого столбца
        counter = 0
        for row in sdl_table.rows:
            for cell in row.cells:
                cell.paragraphs[0].add_run('')
                rc = cell.paragraphs[0].runs[0]
                rc.font.size = Pt(14)
                rc.font.name = 'TimesNewRoman'
                if counter == 0:
                    rc.font.bold = True
                    rc.font.underline = True
                    counter += 1
                else:
                    counter = 0
                set_cell_border(
                    cell,
                    top={"sz": 0},
                    bottom={"sz": 0},
                    start={"sz": 0},
                    end={"sz": 0},
                )

        document.add_paragraph()
        art_header = document.add_paragraph('')
        art_header.paragraph_format.first_line_indent = Inches(0.5)
        art_header.add_run(
            'В США, штат Миссури, город Осейдж-Бич был поврежден плотномер, содержащий цезий-137 и амерций-241')
        art_header.runs[0].font.name = 'TimesNewRoman'
        art_header.runs[0].font.size = Pt(14)
        art_header.runs[0].font.bold = True

        article = document.add_paragraph('')
        article.paragraph_format.first_line_indent = Inches(0.5)
        article.add_run(
            'В 11:46 CDT 7/3/2024 сотрудник по радиационной безопасности в Midwest Subsurface Testing сообщил, что на строительной площадке был поврежден датчик. Датчик плотности влаги InstroTek MC1 Elite, содержащий 10 милликюри цезия-137 и 50 милликюри америция-241/бериллия, был опрокинут погрузчиком. Источник застрял в экранированном положении. Было проведено радиологическое обследование, которое подтвердило отсутствие загрязнения. Поврежденный датчик был извлечен и доставлен на предприятие поставщика для проведения испытания на герметичность.\nЭто событие было сообщено в соответствии с 10 CFR 30.50 (b)(2).')
        article.runs[0].font.name = 'TimesNewRoman'
        article.runs[0].font.size = Pt(14)

        document.save('demo.docx')


class TranslateTablesFrame(tk.Frame):
    def __init__(self, parent):
        super().__init__(parent)
        self.tt_canvas = tk.Canvas(self, width=800, height=600, bg='#6A5ACD')
        self.tt_canvas.create_image((10, 10), anchor='nw', image=back_img, tags='back_but')
        self.tt_canvas.tag_bind('back_but', '<Button-1>', lambda win='menu': goto(event=None, win='menu'))
        self.tt_canvas.pack()
        self.pack()


class MainWin:
    def __init__(self, master):
        global index, frameList
        global panel_img, report_img, table_img, back_img, param_img, info_img, confirm_img, dismiss_img, restore_img
        global restore_kw_img, restore_kw_gif, save_kw_img
        global ru_flag, uk_flag, us_flag, fr_flag, ge_flag, eu_flag

        panel_img = tk.PhotoImage(file='images/panel.png')
        report_img = tk.PhotoImage(file='images/report.png')
        table_img = tk.PhotoImage(file='images/translate.png')
        back_img = tk.PhotoImage(file='images/back36.png')
        param_img = tk.PhotoImage(file='images/param.png')
        info_img = tk.PhotoImage(file='images/info.png')
        confirm_img = tk.PhotoImage(file='images/confirm24.png')
        dismiss_img = tk.PhotoImage(file='images/dismiss24.png')
        restore_img = tk.PhotoImage(file='images/restore.png')
        restore_kw_img = tk.PhotoImage(file='images/restore_kw.png')
        save_kw_img = tk.PhotoImage(file='images/save_kw.png')

        ru_flag = tk.PhotoImage(file='images/rus.png')
        uk_flag = tk.PhotoImage(file='images/uk.png')
        us_flag = tk.PhotoImage(file='images/us.png')
        fr_flag = tk.PhotoImage(file='images/fr.png')
        ge_flag = tk.PhotoImage(file='images/ge.png')
        eu_flag = tk.PhotoImage(file='images/eu.png')

        combobox_style = ttk.Style()
        combobox_style.theme_create('combobox_style', parent='alt',
                                    settings={'TCombobox':
                                                  {'configure':
                                                       {'selectbackground': '#6A5ACD',
                                                        'fieldbackground': '#6A5ACD',
                                                        'background': '#6A5ACD',
                                                        'foreground': '#6A5ACD'
                                                        }}}
                                    )
        combobox_style.theme_use('combobox_style')

        mainframe = tk.Frame(master)
        mainframe.pack()

        index = 0
        frameList = [MenuFrame(mainframe), WebSourcesFrame(mainframe), CompileReportsFrame(mainframe),
                     TranslateTablesFrame(mainframe)]
        for frame in frameList[1:]:
            frame.forget()


def goto(event, win):
    global index
    frameList[index].forget()
    if win == 'menu':
        index = 0
    elif win == 'ws':
        index = 1
    elif win == 'cr':
        index = 2
    elif win == 'tt':
        index = 3
    frameList[index].tkraise()
    frameList[index].pack()


def set_config():
    global browser, browser_path, download_path
    # изначально все = default
    with open('config.txt', 'r') as file:
        config = [line.strip() for line in file.readlines()]

    for elem in config:
        if elem[:10] == 'browser = ':
            browser = elem[10:].strip().lower()
        elif elem[:15] == 'browser_path = ':
            browser_path = elem[15:].strip()
        elif elem[:16] == 'download_path = ':
            download_path = elem[16:].strip()

    if download_path != 'default' and download_path[-1] != '/':
        download_path.join('/')

    print(browser_path, browser)
    if browser != 'default':
        wb.register(browser, None, wb.BackgroundBrowser(browser_path))


with open('smi.txt', 'r') as file:
    # чтобы убрать \n в конце строк использую стрип
    smi_list = [line.strip() for line in file.readlines()]

set_config()
root = tk.Tk()
root.resizable(False, False)
obj = MainWin(root)
root.mainloop()
