"""

"""
import requests
import pprint
import html
import re
import random
import g4f
import time
from deep_translator import GoogleTranslator
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, Inches, Cm
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT


class SingleRepDocx:
    def __init__(self, url):
        """
        Инициализация документа
        """
        self.rep_dict = self.create_reports_dict(url)
        document = Document()

        section = self.document.sections[0]
        header = section.header
        p = header.add_paragraph("Приложение №2")
        p.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

        sections = self.document.sections
        for section in sections:
            section.top_margin = Cm(2.0)
            section.bottom_margin = Cm(2.0)
            section.left_margin = Cm(2.5)
            section.right_margin = Cm(1.0)

        sdl_table = self.document.add_table(rows=3, cols=2)

        column0 = sdl_table.columns[0]
        column0.width = Inches(2.1)
        column1 = sdl_table.columns[1]
        column1.width = Inches(5)

        cell0 = sdl_table.cell(row_idx=0, col_idx=0).text = "Источник:"
        cell1 = sdl_table.cell(row_idx=0, col_idx=1).text = "«nrc.gov»"
        cell2 = sdl_table.cell(row_idx=1, col_idx=0).text = "Опубликовано:"
        cell3 = sdl_table.cell(row_idx=1, col_idx=1).text = list(self.rep_dict.values())[0][0]
        cell4 = sdl_table.cell(row_idx=2, col_idx=0).text = "Ссылка на источник:"
        cell5 = sdl_table.cell(row_idx=2, col_idx=1).text = url

        # обнуление обводки, смена шрифта, размера шрифта + bold

        counter = 0
        for row in sdl_table.rows:
            for cell in row.cells:
                cell.paragraphs[0].add_run('')
                rc = cell.paragraphs[0].runs[0]
                rc.font.size = Pt(14)
                rc.font.name = 'TimesNewRoman'
                if counter == 0:
                    rc.font.bold = True
                    rc.font.underline = True
                    counter += 1
                else:
                    counter = 0
                self.set_cell_border(
                    cell,
                    top={"sz": 0},
                    bottom={"sz": 0},
                    start={"sz": 0},
                    end={"sz": 0},
                )

    def save_doc(self):
        """
        Сохранение документа
        """
        art_header = self.document.add_paragraph('')
        art_header.paragraph_format.first_line_indent = Cm(1.25)
        art_header.paragraph_format.line_spacing = 1.15
        self.art_header_text = list(self.rep_dict.values())[0][1]
        art_header.add_run(self.art_header)
        art_header.runs[0].font.name = 'TimesNewRoman'
        art_header.runs[0].font.size = Pt(14)
        art_header.runs[0].font.bold = True
        art_header.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        article = self.document.add_paragraph('')
        article.paragraph_format.first_line_indent = Cm(1.25)
        article.paragraph_format.line_spacing = 1.15
        article.add_run(list(self.rep_dict.values())[0][2])
        article.runs[0].font.name = 'TimesNewRoman'
        article.runs[0].font.size = Pt(14)
        article.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY # ИЛИ DISTRIBUTE

        self.document.save('demo.docx')

    def show_header(self):
        """
        демонстрирование текста
        крепится к canvas.item_config
        """
        return list(self.rep_dict.values())[0][1]

    def change_header(self):
        """
        изменение заголовка
        меняется элемент словаря, после чего возвращается егозначение
        крепится к canvas.item_config
        """
        ai_list = [g4f.models.gpt_4o, g4f.models.gemeni_pro, g4f.models.claude_3_haiku]
        ai_model = random.choice(ai_list)

        list(self.rep_dict.values())[0][1] = g4f.ChatCompletion.create(
            model = ai_model,
            messages=[{'role': 'user',
                           'content': f'не отвечай ничего лишнего, дай мне краткое содержание текста можно использовать до 15 слов без слов "отчет", "текст": {report}'}]
            )

        self.show_header()

    def set_cell_border(self, cell, **kwargs):
        """
        set_cell_border(
            cell,
            top={"sz": 12, "val": "single", "color": "#FF0000", "space": "0"},
            bottom={"sz": 12, "color": "#00FF00", "val": "single"},
            start={"sz": 24, "val": "dashed", "shadow": "true"},
            end={"sz": 12, "val": "dashed"},
        )
        """
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()

        # check for tag existnace, if none found, then create one
        tcBorders = tcPr.first_child_found_in("w:tcBorders")
        if tcBorders is None:
            tcBorders = OxmlElement('w:tcBorders')
            tcPr.append(tcBorders)

        # list over all available tags
        for edge in ('start', 'top', 'end', 'bottom', 'insideH', 'insideV'):
            edge_data = kwargs.get(edge)
            if edge_data:
                tag = 'w:{}'.format(edge)

                # check for tag existnace, if none found, then create one
                element = tcBorders.find(qn(tag))
                if element is None:
                    element = OxmlElement(tag)
                    tcBorders.append(element)

                # looks like order of attributes is important
                for key in ["sz", "val", "color", "space", "shadow"]:
                    if key in edge_data:
                        element.set(qn('w:{}'.format(key)), str(edge_data[key]))

    
    def create_reports_dict(self, url: str) -> dict[str, list[str]] | str:

        def create_header(report: str) -> str:
            response = g4f.ChatCompletion.create(
                model=g4f.models.gpt_4o_mini,
                messages=[{'role': 'user',
                           'content': f'не отвечай ничего лишнего, дай мне краткое содержание текста можно использовать до 15 слов без слов "отчет", "текст": {report}'}]
            )
            return response

        response = requests.get(url)
        if response.status_code == 200:
            html_content = response.text
            id = url[-5:]
            trash_reports = html_content.split('<div class="border">')[1:]

            # ids
            id_date_report = dict()
            date_block = html_content.split(f'<div class="grid border" id="en{id}">')[1].split('</div>')[3].strip()

            try:
                date = date_block.split(r'<b>Notification Date:</b> ')[1].split('<br>')[0]
                date = '%s.%s.%s' % (date[3:5], date[0:2], date[6:])
            except IndexError:
                date_block = html_content.split(f'<div class="grid border" id="en{id}">')[1].split('</div>')[5].strip()
                date = date_block.split(r'<b>Notification Date:</b> ')[1].split('<br>')[0]
                date = '%s.%s.%s' % (date[3:5], date[0:2], date[6:])

            eng_report = re.sub(r'<[^>]+>', '', html.unescape(trash_reports[0].split('</div>')[0])).strip()

            try:
                eng_report = eng_report.split('EN Revision Text: ')[1]
            except IndexError:
                pass

            rus_report = GoogleTranslator(source='en', target='ru').translate(eng_report)

            header = create_header(rus_report)
            time.sleep(1)

            id_date_report[id] = [date, header, rus_report]

            return id_date_report
        elif response.status_code == 404:
            return 'Происшествий по сегодняшней дате не найдено!'
        else:
            return 'Внутренняя ошибка'


url = f'https://www.nrc.gov/reading-rm/doc-collections/event-status/event/2024/20240719en.html#en57210'
